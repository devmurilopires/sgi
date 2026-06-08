import os
from datetime import datetime
from docx import Document
from src.modulos.projetos_mobilidade.parecer.repository import ParecerProjetosMobilidadeRepository
from config.settings import RAIZ_REDE

try:
    from src.core.shared.utils import resource_path
except ImportError:
    def resource_path(path): return path

class ParecerProjetosMobilidadeService:
    def __init__(self):
        self.repo = ParecerProjetosMobilidadeRepository()

    def processar_geracao_parecer(self, dados_form, usuario_logado):
        tipo = dados_form.get('tipo', 'DEFERIDO').upper()
        origem = dados_form.get('origem', '').strip()
        processo = dados_form.get('processo', '').upper()
        assunto = dados_form.get('assunto', '').upper()
        solicitante = dados_form.get('solicitante', '').upper()
        
        if not usuario_logado:
            return False, "Erro de autenticação: Usuário não identificado na sessão."
        criado_por = f"%{usuario_logado}%"

        # Obtém número de forma automática
        numero_parecer = self.repo.obter_proximo_numero_parecer()
        ano_atual = datetime.now().year
        numero_parecer_ano = f"{numero_parecer:03d}/{ano_atual}"

        if not os.path.exists(RAIZ_REDE):
            return False, f"Parecer {numero_parecer_ano} salvo no banco, mas a rede está inacessível para gerar o documento."

        pasta_base = rf"{RAIZ_REDE}\PROJETOS DE MOBILIDADE\{ano_atual}\PARECERES\{tipo}"
        nome_pasta = f"PARECER {numero_parecer_ano.replace('/', '-')} - PROCESSO {processo.replace('/', '-')}"
        caminho_pasta = os.path.join(pasta_base, nome_pasta)
        nome_arquivo = f"PARECER {numero_parecer_ano.replace('/', '-')} - PROJETOS DE MOBILIDADE.docx"
        destino_docx = os.path.join(caminho_pasta, nome_arquivo)

        dados_db = {
            "numero_parecer": numero_parecer,
            "tipo": tipo,
            "processo": processo,
            "origem": origem,
            "assunto": assunto,
            "solicitante": solicitante,
            "caminho_arquivo": destino_docx,
            "criado_por": criado_por
        }

        sucesso_db, resultado_db = self.repo.salvar_parecer_no_banco(dados_db)
        if not sucesso_db:
            return False, f"Falha ao registrar no banco de dados:\n{resultado_db}"

        modelo_str = "dados/modelo_parecer_deferido_pm.docx" if tipo == "DEFERIDO" else "dados/modelo_parecer_indeferido_pm.docx"
        caminho_modelo = resource_path(modelo_str)

        try:
            os.makedirs(caminho_pasta, exist_ok=True)
            self._gerar_documento(caminho_modelo, destino_docx, numero_parecer_ano, processo, assunto, solicitante, datetime.now().strftime('%d/%m/%Y'))
            return True, f"Parecer Técnico Nº {numero_parecer_ano} gerado com sucesso!\nSalvo em:\n{destino_docx}"
        except Exception as e:
            return False, f"Parecer registrado no banco, mas falhou ao criar o Word:\n{e}"

    def _substituir_texto_com_runs(self, paragrafo, mapeamento):
        """
        Substitui as tags de forma segura mantendo a formatação exata de cada fragmento (run).
        Isso impede que trechos normais fiquem em negrito sem querer.
        """
        texto_completo = "".join(run.text for run in paragrafo.runs)
        
        possui_tag = any(chave in texto_completo for chave in mapeamento)
        if not possui_tag:
            return

        for chave, valor in mapeamento.items():
            if chave in texto_completo:
                for run in paragrafo.runs:
                    if chave in run.text:
                        run.text = run.text.replace(chave, valor)
                
                texto_completo = "".join(run.text for run in paragrafo.runs)
                if chave in texto_completo:
                    texto_completo = texto_completo.replace(chave, valor)
                    for i, run in enumerate(paragrafo.runs):
                        if i == 0:
                            run.text = texto_completo
                        else:
                            run.text = ""

    def _gerar_documento(self, modelo_path, destino_path, num_parecer, processo, assunto, solicitante, data_str ):
        doc = Document(modelo_path)
        mapeamento = {
            "{{NUM_PARECER}}": num_parecer,
            "{{PROCESSO}}": processo,
            "{{ASSUNTO}}": assunto,
            "{{SOLICITANTE}}": solicitante,
            "{{DATA}}": data_str
        }
        
        for paragrafo in doc.paragraphs:
            self._substituir_texto_com_runs(paragrafo, mapeamento)
                
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        self._substituir_texto_com_runs(paragrafo, mapeamento)
                
        doc.save(destino_path)