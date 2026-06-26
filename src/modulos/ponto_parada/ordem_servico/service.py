import os
import unicodedata
from datetime import datetime
from docx import Document
import shutil
from docx.shared import Inches
from src.modulos.ponto_parada.ordem_servico.repository import OSRepository
from config.settings import RAIZ_REDE

try:
    from src.core.shared.utils import resource_path
except ImportError:
    try:
        from src.core.shared.utils import resource_path
    except ImportError:
        def resource_path(path): return path

class OSService:
    def __init__(self):
        self.repo = OSRepository()

    # =========================================================
    # REGRAS DE NEGÓCIO E VALIDAÇÕES
    # =========================================================
    def normalizar(self, texto: str) -> str:
        if not texto: return ""
        nfkd = unicodedata.normalize('NFKD', texto)
        sem_acentos = "".join(c for c in nfkd if not unicodedata.combining(c))
        return sem_acentos.upper()

    def consultar_endereco(self, id_procurado):
        return self.repo.buscar_endereco_por_id(id_procurado)

    def obter_historico_formatado(self, id_procurado):
        historico = self.repo.buscar_historico_os(id_procurado)
        if not historico:
            return "Nenhuma movimentação de Ordem de Serviço encontrada para este ID."
        
        texto = ""
        for r in historico:
            texto += f"OS: {r[0]} | Data: {r[1]} | Tipo: {r[2]} | Item: {r[3]}\nEndereço: {r[4]} - Bairro: {r[5]}\nCriado por: {r[6]}\n"
            texto += "-" * 50 + "\n"
        return texto

    # NOVA FUNÇÃO: Retorna os itens correspondentes ao modelo selecionado
    def obter_itens_por_modelo(self, modelo):
        contexto = 'ITEM_URBMIDIA' if modelo == 'URBMÍDIA' else 'ITEM_MCMENSAGEM'
        return self.repo.buscar_tipos_por_contexto(contexto)

    # =========================================================
    # ORQUESTRAÇÃO PRINCIPAL (GERAÇÃO SEGURA - DB PRIMEIRO)
    # =========================================================
    def processar_criacao_os(self, descricoes_acumuladas, modelo_operacao, modelo_escolhido, tipo_os, tipo_item, processo, usuario_logado, origem_demanda):
        if not descricoes_acumuladas:
            return False, "Adicione pelo menos um item (descrição) na lista antes de gerar a OS."

        ano_atual = datetime.now().strftime('%Y')
        
        if not os.path.exists(RAIZ_REDE):
            return False, f"A raiz da rede não está acessível no momento. Verifique a ligação:\n{RAIZ_REDE}"

        # Direciona para as pastas corretas
        if "MCMENSAGEM" in str(modelo_operacao).upper().replace(" ", ""):
            pasta_base = rf"{RAIZ_REDE}\PONTO DE PARADA\{ano_atual}\ORDENS DE SERVICO\MC MENSAGEM"
            item_contexto = "ITEM_MCMENSAGEM"
        else:
            pasta_base = rf"{RAIZ_REDE}\PONTO DE PARADA\{ano_atual}\ORDENS DE SERVICO\URBMIDIA"
            item_contexto = "ITEM_URBMIDIA"

        # Extrai os IDs para relacionamento na Base de Dados
        ids_unicos = list(set([d["id"] for d in descricoes_acumuladas]))
        id_principal = descricoes_acumuladas[0]["id"]
        pontos_adicionais = [pid for pid in ids_unicos if pid != id_principal]

        #  Passamos o modelo_operacao para o repositório separar a numeração!
        numero_os = self.repo.obter_proximo_numero_os(ano_atual, modelo_operacao)
        
        data_str = datetime.now().strftime("%d/%m/%Y")

        tipo_os_up = str(tipo_os).strip().upper() if tipo_os else ""
        tipo_item_up = str(tipo_item).strip().upper() if tipo_item else ""

        # =====================================================================
        # BLINDAGEM SÊNIOR: Encontrar o número definitivo validando as pastas físicas
        # =====================================================================
        nome_pasta = f"{numero_os:03d}-{datetime.now().strftime('%m')}-{ano_atual}-ID{'-'.join(ids_unicos) if ids_unicos else 'EMERGENCIA'}"
        caminho_pasta = os.path.join(pasta_base, nome_pasta)
        nome_arquivo = f"O.S {numero_os:03d}-{ano_atual}-ID{'-'.join(ids_unicos) if ids_unicos else 'EMERGENCIA'}.docx"
        destino_docx = os.path.join(caminho_pasta, nome_arquivo)

        # Verifica se a pasta ou o ficheiro já existem. Se sim, salta para o próximo número!
        while os.path.exists(caminho_pasta) or os.path.exists(destino_docx):
            numero_os += 1
            nome_pasta = f"{numero_os:03d}-{datetime.now().strftime('%m')}-{ano_atual}-ID{'-'.join(ids_unicos) if ids_unicos else 'EMERGENCIA'}"
            caminho_pasta = os.path.join(pasta_base, nome_pasta)
            nome_arquivo = f"O.S {numero_os:03d}-{ano_atual}-ID{'-'.join(ids_unicos) if ids_unicos else 'EMERGENCIA'}.docx"
            destino_docx = os.path.join(caminho_pasta, nome_arquivo)
        # =====================================================================

        # Monta os dados para a Base de Dados
        dados_db = {
            "processo": processo,
            "numero_os": numero_os,
            "data_criacao": datetime.strptime(data_str, "%d/%m/%Y").date(),
            "id_principal": id_principal,
            "origem": origem_demanda,
            "acao": tipo_os_up.upper(),
            "item": tipo_item_up,
            "item_contexto": item_contexto,
            "modelo": modelo_operacao,
            "descricao": "\n".join([item["descricao"] for item in descricoes_acumuladas]),
            "usuario": f"%{usuario_logado}%",
            "caminho": destino_docx,
            "pontos_adicionais": pontos_adicionais
        }

        # 1º PASSO: Tenta guardar na Base de Dados (Se falhar aqui, não mexe na rede física e para o processo!)
        try:
            self.repo.salvar_os(dados_db)
        except Exception as e:
            return False, f"Erro Crítico! A OS NÃO foi gerada pois houve falha na Base de Dados:\n{str(e)}"

        # 2º PASSO: Tenta criar a pasta e o Ficheiro Word
        try:
            os.makedirs(caminho_pasta, exist_ok=True)
            caminho_modelo = resource_path(modelo_escolhido)
            self._gerar_documento_modelo(caminho_modelo, destino_docx, numero_os, data_str, id_principal, descricoes_acumuladas, processo)
            return True, f"Ordem de Serviço Nº {numero_os:03d} criada e registada com sucesso!\nSalva em:\n{destino_docx}"
            
        except Exception as e:
            # 3º PASSO (ROLLBACK FÍSICO): Se a geração do Word falhar (ex: rede foi abaixo a meio), apaga a pasta que acabou de criar!
            if os.path.exists(caminho_pasta):
                try:
                    shutil.rmtree(caminho_pasta)
                except Exception:
                    pass
            return False, f"Atenção: A OS foi registada na base de dados, mas houve falha ao gerar o documento Word na rede:\n{e}"
                
    def _gerar_documento_modelo(self, modelo_path, destino_path, numero_os, data_str, id_texto, descricoes, processo_str):
        doc = Document(modelo_path)
        # Adicionada a tag {{PROCESSO}} para que saia automaticamente no template do Word
        mapeamento = {
            "{{NUMERO_OS}}": f"{numero_os:03d}",
            "{{DATA}}": data_str,
            "{{ID}}": id_texto if id_texto.strip() else "-",
            "{{PROCESSO}}": processo_str
        }
        
        for paragrafo in doc.paragraphs:
            texto_original = "".join(run.text for run in paragrafo.runs)
            novo_texto = texto_original
            for chave, valor in mapeamento.items():
                if chave in novo_texto:
                    novo_texto = novo_texto.replace(chave, valor)
            if novo_texto != texto_original:
                for run in paragrafo.runs: run.text = ""
                if paragrafo.runs: paragrafo.runs[0].text = novo_texto

        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        texto_original = "".join(run.text for run in paragrafo.runs)
                        novo_texto = texto_original
                        for chave, valor in mapeamento.items():
                            if chave in novo_texto:
                                novo_texto = novo_texto.replace(chave, valor)
                        if novo_texto != texto_original:
                            for run in paragrafo.runs: run.text = ""
                            if paragrafo.runs: paragrafo.runs[0].text = novo_texto

        for paragrafo in doc.paragraphs:
            if "{{DESCRICAO}}" in paragrafo.text:
                p = paragrafo._element
                parent = p.getparent()
                
                tabela = doc.add_table(rows=1, cols=2)
                tabela.style = 'Table Grid'
                tabela.columns[0].width = Inches(1.0)
                tabela.columns[1].width = Inches(5.0)
                
                hdr_cells = tabela.rows[0].cells
                hdr_cells[0].text = 'ID'
                hdr_cells[1].text = 'DESCRIÇÃO'
                
                for item in descricoes:
                    row_cells = tabela.add_row().cells
                    row_cells[0].text = item['id']
                    row_cells[1].text = item['descricao']
                
                p.addnext(tabela._element)
                parent.remove(p)
                break
        
        doc.save(destino_path)