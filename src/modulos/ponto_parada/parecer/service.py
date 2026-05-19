import os
import re
from datetime import datetime
from docx import Document
from src.modulos.ponto_parada.parecer.repository import ParecerRepository
from config.settings import RAIZ_REDE

try:
    from src.core.shared.utils import resource_path
except ImportError:
    def resource_path(path): return path

class ParecerService:
    def __init__(self):
        self.repo = ParecerRepository()

    # NOVAS FUNÇÕES: Buscam dinamicamente as listas para popular a interface
    def obter_assuntos(self):
        return self.repo.buscar_parametro_generico('ASSUNTO_PARECER')

    def obter_solicitantes(self):
        return self.repo.buscar_parametro_generico('SOLICITANTE_PARECER')
    
    def obter_itens(self):
        return self.repo.buscar_todos_itens()

    def _converter_quantidade(self, qtd_str):
        if not qtd_str or qtd_str == "-": return None
        nums = re.findall(r'\d+', qtd_str)
        if nums: return float(nums[0])
        map_qtd = {"um": 1, "uma": 1, "dois": 2, "duas": 2, "tres": 3, "três": 3, "quatro": 4, "cinco": 5}
        return map_qtd.get(qtd_str.lower().strip(), None)

    def processar_geracao_parecer(self, dados_form, ids_list, usuario_logado):
        if not ids_list:
            return False, "Adicione ao menos um ID antes de gerar o parecer."

        origem = dados_form['origem']
        tipo_parecer = dados_form['tipo']
        processo = dados_form['processo'] or "-"
        assunto = dados_form['assunto'] or "-"
        solicitante = dados_form['solicitante'] or "-"
        tipo_exec = dados_form['tipo_execucao'] or "-"
        item = dados_form['item'] or "-"
        endereco = dados_form['endereco'] or "-"
        motivo = dados_form['motivo'] if tipo_parecer.upper() == "INDEFERIDO" else None
        quantidade_texto = dados_form['quantidade'] or "-"
        
        ids_joined = ", ".join(ids_list)
        data_atual = datetime.now()
        ano = data_atual.year
        data_str = data_atual.strftime("%d/%m/%Y")

        quantidade_normalizada = quantidade_texto.lower().strip()
        plurais = {
            "Abrigo Metálico": "Abrigos Metálicos",
            "Placa/Barrote": "Placas/Barrote",
            "Placa/Poste": "Placas/Poste",
            "Parada Segura": "Paradas Seguras",
            "Abrigo Concreto": "Abrigos Concretos"
        }
        if not (quantidade_normalizada.startswith("um") or quantidade_normalizada.startswith("uma")):
            item = plurais.get(item, item)

        try:
            numero = self.repo.obter_proximo_numero(ano)
        except Exception as e:
            return False, str(e)

        modelo = resource_path(os.path.join("dados", "modelo_deferido_pp.docx")) if tipo_parecer == "Deferido" else resource_path(os.path.join("dados", "modelo_indeferido_pp.docx"))
        
        if not os.path.exists(modelo):
            return False, f"Modelo Word não encontrado em: {modelo}"
        
        if not os.path.exists(RAIZ_REDE):
            return False, f"A raiz da rede não está acessível no momento. Verifique a conexão:\n{RAIZ_REDE}"

        pasta_base = rf"{RAIZ_REDE}\PONTO DE PARADA\{ano}\PARECERES TECNICOS"
        pasta_saida = os.path.join(pasta_base, tipo_parecer.upper())
        nome_arquivo = f"Parecer_{numero:03d}_{ano}_{tipo_parecer}.docx"
        caminho_arquivo = os.path.join(pasta_saida, nome_arquivo)

        dados_db = {
            "numero": numero,
            "ano": ano,
            "tipo_parecer": tipo_parecer,
            "processo": processo,
            "assunto": assunto,
            "solicitante": solicitante,
            "tipo_exec": tipo_exec,
            "item": item,
            "endereco": endereco,
            "quantidade": self._converter_quantidade(quantidade_texto),
            "motivo": motivo,
            "caminho_arquivo": caminho_arquivo,
            "usuario_logado": f"%{usuario_logado}%",
            "origem": origem,
            "ids_list": ids_list
        }

        try:
            self.repo.salvar_parecer(dados_db)
        except Exception as e:
            return False, f"Erro Crítico! O Parecer NÃO foi gerado pois houve falha no Banco de Dados:\n{str(e)}"

        try:
            os.makedirs(pasta_saida, exist_ok=True)
            self._gerar_documento_word(modelo, caminho_arquivo, {
                "{{NUM_PARECER}}": f"{numero:03d}",
                "{{DATA}}": data_str,
                "{{PROCESSO}}": processo,
                "{{ASSUNTO}}": assunto,
                "{{SOLICITANTE}}": solicitante,
                "{{ID}}": ids_joined,
                "{{TIPO}}": tipo_exec,
                "{{ITEM}}": item,
                "{{ENDERECO}}": endereco,
                "{{MOTIVO}}": motivo or "-",
                "{{QUANTIDADE}}": quantidade_texto
            })
            return True, f"Parecer {numero:03d}/{ano} criado com sucesso!\nSalvo em:\n{caminho_arquivo}"
            
        except Exception as e:
            return False, f"Atenção: O Parecer foi registrado no banco, mas falhou ao gerar o Word:\n{e}"

    def _gerar_documento_word(self, modelo_path, destino_path, tags):
        doc = Document(modelo_path)
        for p in doc.paragraphs:
            texto_completo = "".join(run.text for run in p.runs)
            modificado = False
            for tag, val in tags.items():
                if tag in texto_completo:
                    texto_completo = texto_completo.replace(tag, val)
                    modificado = True
            if modificado:
                for run in p.runs: run.text = ""
                if p.runs: p.runs[0].text = texto_completo
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        texto_completo = "".join(run.text for run in p.runs)
                        modificado = False
                        for tag, val in tags.items():
                            if tag in texto_completo:
                                texto_completo = texto_completo.replace(tag, val)
                                modificado = True
                        if modificado:
                            for run in p.runs: run.text = ""
                            if p.runs: p.runs[0].text = texto_completo
        doc.save(destino_path)