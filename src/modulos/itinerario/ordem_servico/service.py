import os
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from src.core.shared.utils import resource_path
from src.modulos.itinerario.ordem_servico.repository import OSItinerarioRepository

class OSItinerarioService:
    def __init__(self):
        self.repo = OSItinerarioRepository()
        self.ano_atual = datetime.now().year
        self.output_root = r"\\172.20.0.57\dados\DIPLA\ARQUIVOS SIGP - SIGA - SPR\ITINERARIO\ORDENS DE SERVICO"
        os.makedirs(self.output_root, exist_ok=True)

    def formatar_lista_com_e(self, lista):
        lst = [s for s in lista if s]
        if not lst: return ""
        if len(lst) == 1: return lst[0]
        if len(lst) == 2: return f"{lst[0]} e {lst[1]}"
        return ", ".join(lst[:-1]) + f" e {lst[-1]}"

    def buscar_sugestoes(self, tipo):
        if tipo == "EMPRESAS": return self.repo.buscar_empresas()
        if tipo == "LINHAS": return self.repo.buscar_linhas()
        return []

    def _replace_tags_in_doc(self, doc, mapping):
        for para in doc.paragraphs:
            for key, val in mapping.items():
                if key in para.text:
                    for r in para.runs:
                        if key in r.text: r.text = r.text.replace(key, str(val))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, val in mapping.items():
                            if key in para.text:
                                for r in para.runs:
                                    if key in r.text: r.text = r.text.replace(key, str(val))

    def _inserir_anexos(self, doc, tag, anexos, legenda_path=None):
        target_para = next((p for p in doc.paragraphs if tag in p.text), None)
        if not target_para and not anexos: return
        
        for idx, anexo in enumerate(anexos):
            p = target_para if (idx == 0 and target_para) else doc.add_paragraph()
            if idx == 0 and target_para: p.text = "" 

            ida, volta = self.formatar_lista_com_e(anexo.get("linhas_ida", [])), self.formatar_lista_com_e(anexo.get("linhas_volta", []))
            r_ida, r_volta = anexo.get("ruas_ida", ""), anexo.get("ruas_volta", "")
            img_path = anexo.get("image_path")
            
            if ida or r_ida: doc.add_paragraph().add_run(f"Linha {ida}, quando trafegar no sentido Ida: {r_ida}.")
            if volta or r_volta: doc.add_paragraph().add_run(f"Linha {volta}, quando trafegar no sentido Volta: {r_volta}.")
            
            if img_path and os.path.exists(img_path):
                try: doc.add_paragraph().add_run().add_picture(img_path, width=Inches(6.0))
                except: pass
                
            if legenda_path and os.path.exists(legenda_path):
                try: doc.add_paragraph().add_run().add_picture(legenda_path, width=Inches(2.0))
                except: pass
            doc.add_paragraph("")

    def processar_criacao_os(self, tipo_os, form_dados, empresas, linhas, anexos_raw, usuario):
        num_os = self.repo.obter_proximo_numero_os(self.output_root)
        num_os_str = f"{num_os:03d}"
        
        tmpl_nome = f"modelo_os_{tipo_os.lower()}.docx"
        tmpl_path = resource_path(f"dados/{tmpl_nome}")
        legenda_path = resource_path(f"dados/img_legenda_{'obra' if tipo_os == 'OBRAS' else 'evento'}.png")

        if not os.path.exists(tmpl_path):
            return False, f"O modelo {tmpl_nome} não foi encontrado na pasta 'dados'."

        # TRATAMENTO INTELIGENTE DA DATA NO WORD
        data_text = ""
        datas_raw = form_dados.get("datas", [])
        modo_data = form_dados.get("modo_data", "PERIODO")

        if datas_raw:
            if modo_data == "ISOLADOS":
                data_text = "nos dias " + self.formatar_lista_com_e(datas_raw) if len(datas_raw) > 1 else f"no dia {datas_raw[0]}"
            else:
                data_text = f"no dia {datas_raw[0]}" if len(datas_raw) == 1 else f"no período de {datas_raw[0]} a {datas_raw[1]}"

        mapping = {
            "{{NUM_OS}}": num_os_str,
            "{{DATA}}": datetime.now().strftime("%d/%m/%Y"),
            "{{N_PROCESSO}}": form_dados.get("processo", ""),
            "{{EMPRESAS}}": "\n".join(empresas),
            "{{EMPRESA}}": "\n".join(empresas),
            "{{LINHAS}}": " " + self.formatar_lista_com_e(linhas) if linhas else "",
            "{{HORARIO_INICIO}}": form_dados.get("hr_inicio", ""),
            "{{HORARIO_FINAL}}": form_dados.get("hr_fim", ""),
            "{{ANEXO}}": "", "{{IMG_LEGENDA}}": ""
        }

        if tipo_os == "EVENTOS":
            mapping.update({"{{ENDERECO}}": form_dados.get("endereco", ""), "{{EVENTO}}": form_dados.get("evento", ""), "{{DATA_EVENTO}}": data_text, "{{LINHA_ESP}}": "", "{{RUAS_IDA}}": "", "{{RUAS_VOLTA}}": "", "{{NUM_PAGINA}}": ""})
        elif tipo_os == "CORRIDA":
            mapping.update({"{{NOME_CORRIDA}}": form_dados.get("nome_corrida", ""), "{{DATA_CORRIDA}}": data_text, "{{KM}}": form_dados.get("km", ""), "{{SOLICITANTE}}": form_dados.get("solicitante", ""), "{{HORARIO_INICIAL}}": form_dados.get("hr_inicio", "")})
        elif tipo_os == "OBRAS":
            mapping.update({"{{ENDERECO}}": form_dados.get("endereco", ""), "{{TIPO_OBRA}}": form_dados.get("tipo_obra", ""), "{{DATA_OBRA}}": data_text})

        try:
            doc = Document(tmpl_path)
            self._replace_tags_in_doc(doc, mapping)
            self._inserir_anexos(doc, "{{ANEXO}}", anexos_raw, legenda_path if tipo_os != 'CORRIDA' else None)
            
            filename = f"OS Nº{num_os_str} de {datetime.now().strftime('%d-%m-%Y')} {tipo_os.upper()}.docx"
            filepath = os.path.join(self.output_root, filename)
            while os.path.exists(filepath):
                num_os += 1
                num_os_str = f"{num_os:03d}"
                filename = f"OS Nº{num_os_str} de {datetime.now().strftime('%d-%m-%Y')} {tipo_os.upper()}.docx"
                filepath = os.path.join(self.output_root, filename)
            
            doc.save(filepath)

            # Adicionando a "origem" no pacote de dados pro BD
            dados_db = {
                "num_os": num_os, "ano": self.ano_atual, "tipo": tipo_os.lower(), "processo": form_dados.get("processo", ""),
                "origem": form_dados.get("origem", ""), # <-- NOVO CAMPO
                "empresa_principal": empresas[0] if empresas else None, "empresas_text": "; ".join(empresas), "endereco": form_dados.get("endereco", ""),
                "data_evento_text": data_text if tipo_os == 'EVENTOS' else "", "horario_inicio": form_dados.get("hr_inicio", ""), "horario_final": form_dados.get("hr_fim", ""),
                "linhas_text": self.formatar_lista_com_e(linhas), "linhas_adicionadas_text": ", ".join(linhas), "linhas_especificas_text": "",
                "ruas_ida": "", "ruas_volta": "", "evento": form_dados.get("evento", "") if tipo_os == 'EVENTOS' else form_dados.get("tipo_obra", ""),
                "anexo_filename": os.path.basename(anexos_raw[0].get("image_path", "")) if anexos_raw and anexos_raw[0].get("image_path") else None,
                "pasta_path": self.output_root, "docx_path": filepath, "criado_por": usuario,
                "anexos_json": json.dumps(anexos_raw, ensure_ascii=False),
                "nome_corrida": form_dados.get("nome_corrida", ""), "km": form_dados.get("km", ""), "solicitante": form_dados.get("solicitante", ""),
                "tipo_obra": form_dados.get("tipo_obra", ""), "data_obra_text": data_text if tipo_os == 'OBRAS' else ""
            }
            sucesso, msg_db = self.repo.salvar_os_itinerario(dados_db)
            if not sucesso: return False, msg_db

            return True, f"OS {num_os_str} gerada com sucesso em {self.output_root}!"
        except Exception as e:
            return False, f"Erro interno ao gerar o arquivo Word: {str(e)}"